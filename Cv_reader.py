from groq import Groq
from dotenv import load_dotenv
from pathlib import Path
from io import BytesIO
import requests
import os
import json
import docx
from PyPDF2 import PdfReader
from streamlit import audio_input

load_dotenv()

def read_file(text_file_path):
    context_path = Path(__file__).parent / text_file_path
    with open (context_path, "r", encoding="utf-8") as file:
        return file.read()
    
def extract_text_from_pdf(pdf_path):
    reader = PdfReader(pdf_path)
    text = ""
    for page in reader.pages:
        text += page.extract_text() or ""
    return text

def extract_text_from_docx(docx_path):
    doc = docx.Document(docx_path)
    return "\n".join([para.text for para in doc.paragraphs])

def read_cv(file_input):
    client = Groq(api_key=os.environ["GROQ_API_KEY"])
    # Si file_input est un chemin (str ou Path)
    if isinstance(file_input, (str, Path)):
        ext = str(file_input).lower()
        if ext.endswith(".pdf"):
            text = extract_text_from_pdf(file_input)
        elif ext.endswith(".docx"):
            text = extract_text_from_docx(file_input)
        else:
            raise ValueError("Le fichier doit être au format PDF ou DOCX.")
    else:
        # Cas UploadedFile (Streamlit)
        name = file_input.name.lower()
        if name.endswith(".pdf"):
            reader = PdfReader(file_input)
            text = ""
            for page in reader.pages:
                text += page.extract_text() or ""
        elif name.endswith(".docx"):
            doc = docx.Document(file_input)
            text = "\n".join([para.text for para in doc.paragraphs])
        else:
            raise ValueError("Le fichier doit être au format PDF ou DOCX.")

    response = client.chat.completions.create(
        model="openai/gpt-oss-20b",
        messages=[
            {
                "role": "system",
                "content": read_file(text_file_path="context.txt"),
                
            },
            {
                "role": "user",
                "content": f"Analyse le texte ci-dessous (ta réponse doit être dans le format JSON) : {text}",
            }
        ],
        temperature=0.0,
        response_format={"type": "json_object",}
    )
    return response.choices[0].message.content

if __name__ == "__main__":
    # Exemple d'utilisation
    cv_path = "C:\\Users\\sebas\\OneDrive\\Bureau\\M5-HETIC\\Agents\\Projet Scrappy Offres\\job-scrape-engine\\CV Sebastian Data.pdf"  # Remplacez par le chemin de votre CV
    analysis = read_cv(cv_path)
    print(analysis)
